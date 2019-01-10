from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches

#打开文档
document = Document()
#加入不同等级的标题
document.add_heading('Document Title',0)
document.add_heading(u'二级标题',1)
document.add_heading(u'二级标题',2) #添加文本
paragraph = document.add_paragraph(u'添加了文本')
#保存文件
document.save('./data/demo.docx')